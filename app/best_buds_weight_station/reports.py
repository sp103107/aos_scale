"""Non-authoritative harvest reports (JSON / CSV / XLSX / DOCX).

Authoritative truth remains session JSONL + individual record files.
Reports are operator handoff derivatives compiled from accepted weight records.

BBWS SR5 polishes handoff presentation (Book Spine / Output Response cite-only).
Does not claim Metrc compliance or legal-for-trade certification.
"""

from __future__ import annotations

import csv
import hashlib
import re
from collections import defaultdict
from pathlib import Path
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from .spreadsheet import HEADERS, row_for
from .storage import atomic_json, canonical, parse_jsonl

NON_CLAIMS = [
    "Non-authoritative operator handoff report.",
    "Authoritative truth remains session JSONL and individual record files.",
    "Not legal-for-trade / metrology certification.",
    "Not Metrc compliance.",
]

_HEADER_FILL = PatternFill("solid", fgColor="1B6B52")
_HEADER_FONT = Font(bold=True, color="FFFFFF")
_HEADER_ALIGN = Alignment(horizontal="left", vertical="center", wrap_text=True)


def _accepted_rows(session_dir: Path) -> list[dict[str, Any]]:
    rows = [
        row
        for row in parse_jsonl(session_dir / "records.jsonl")
        if row.get("event_type") == "weight_record" and row.get("record_status") == "accepted"
    ]
    rows.sort(key=lambda row: row["sequence"])
    return rows


def plain_export_stem(run_id: str | None, session_id: str) -> str:
    """Windows-safe plain filename stem (no spaces-as-noise; keep underscores)."""
    raw = (run_id or session_id or "harvest").strip()
    cleaned = re.sub(r"[^\w.\-]+", "_", raw, flags=re.UNICODE)
    cleaned = cleaned.strip("._") or "harvest"
    return cleaned[:80]


def _cultivator_for_strain(rows: list[dict[str, Any]], strain: str) -> str:
    for row in rows:
        if str(row.get("cultivar_normalized_name") or "") == strain:
            return str(row.get("facility_id") or row.get("cultivator") or "")
    return ""


def _style_header_row(ws, column_count: int) -> None:
    for col in range(1, column_count + 1):
        cell = ws.cell(1, col)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = _HEADER_ALIGN
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(column_count)}1"


def _autosize_columns(ws, max_width: int = 36) -> None:
    for idx, column_cells in enumerate(ws.columns, start=1):
        length = 0
        for cell in column_cells:
            value = "" if cell.value is None else str(cell.value)
            length = max(length, len(value))
        ws.column_dimensions[get_column_letter(idx)].width = min(max(length + 2, 10), max_width)


def _write_docx(path: Path, report: dict[str, Any], rows: list[dict[str, Any]]) -> None:
    """Write a polished operator harvest summary Word document."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:  # pragma: no cover - exercised when optional dep missing
        raise RuntimeError(
            "python-docx is required for DOCX reports. Install with: pip install python-docx"
        ) from exc

    doc = Document()
    title = doc.add_heading("Best Buds Harvest Weight Report", level=0)
    title.runs[0].font.size = Pt(18)

    doc.add_heading("Cover", level=1)
    cover = doc.add_paragraph()
    cover.add_run("Report ID: ").bold = True
    cover.add_run(f"{report['report_id']}\n")
    cover.add_run("Session ID: ").bold = True
    cover.add_run(f"{report['session_id']}\n")
    cover.add_run("Run ID: ").bold = True
    cover.add_run(f"{report.get('run_id', '')}\n")
    cover.add_run("Cultivator: ").bold = True
    cover.add_run(f"{report.get('cultivator') or '—'}\n")
    cover.add_run("Operator: ").bold = True
    cover.add_run(f"{report.get('operator_id') or '—'}\n")
    cover.add_run("Compiled at: ").bold = True
    cover.add_run(f"{report['compiled_at']}\n")
    cover.add_run("Record count: ").bold = True
    cover.add_run(f"{report['record_count']}\n")
    cover.add_run("Total net weight: ").bold = True
    cover.add_run(f"{report['total_net_g']} g")

    doc.add_paragraph(
        "Non-claim: This export is a non-authoritative operator handoff. "
        "Authoritative truth is the local session JSONL ledger. "
        "Not legal-for-trade or Metrc compliance."
    )

    doc.add_heading("Strain totals", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Cultivator"
    header[1].text = "Strain"
    header[2].text = "Net g"
    for cultivar, net in report["cultivar_totals"].items():
        cells = table.add_row().cells
        cells[0].text = str(report.get("cultivator") or _cultivator_for_strain(rows, cultivar) or "—")
        cells[1].text = str(cultivar)
        cells[2].text = f"{net}"
    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = "TOTAL"
    total_row[2].text = f"{report['total_net_g']}"

    doc.add_heading("Plant records", level=1)
    detail = doc.add_table(rows=1, cols=7)
    detail.style = "Table Grid"
    for idx, name in enumerate(
        ("Seq", "Barcode", "Cultivator", "Strain", "Gross g", "Tare g", "Net g")
    ):
        detail.rows[0].cells[idx].text = name
    for row in rows:
        cells = detail.add_row().cells
        cells[0].text = str(row.get("sequence", ""))
        cells[1].text = str(row.get("barcode_raw", ""))
        cells[2].text = str(row.get("facility_id") or row.get("cultivator") or "")
        cells[3].text = str(row.get("cultivar_normalized_name", ""))
        cells[4].text = str(row.get("gross_g", ""))
        cells[5].text = str(row.get("tare_g", ""))
        cells[6].text = str(row.get("net_g", ""))

    footer = doc.add_paragraph()
    footer.add_run("\nHandoff footer — non-claims: ").bold = True
    footer.add_run(" ".join(report.get("non_claims") or NON_CLAIMS))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def _write_xlsx(
    path: Path,
    report: dict[str, Any],
    rows: list[dict[str, Any]],
) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Summary"
    sheet.append(["Cultivator", "Strain", "Net g"])
    for cultivar, net in report["cultivar_totals"].items():
        sheet.append(
            [
                report.get("cultivator") or _cultivator_for_strain(rows, cultivar) or "",
                cultivar,
                net,
            ]
        )
    sheet.append(["", "TOTAL", report["total_net_g"]])
    _style_header_row(sheet, 3)
    _autosize_columns(sheet)

    detail = workbook.create_sheet("Records")
    detail.append(list(HEADERS))
    for row in rows:
        detail.append(row_for(row))
    _style_header_row(detail, len(HEADERS))
    _autosize_columns(detail, max_width=28)

    claims = workbook.create_sheet("NonClaims")
    claims.append(["Key", "Statement"])
    for idx, line in enumerate(report.get("non_claims") or NON_CLAIMS, start=1):
        claims.append([f"non_claim_{idx}", line])
    claims.append(["authoritative_source", "session records.jsonl"])
    claims.append(["handoff_polish_series", "BBWS_SR5_run_artifact_polish"])
    _style_header_row(claims, 2)
    _autosize_columns(claims, max_width=80)

    path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(path)


def compile_report(session_dir: str | Path) -> dict[str, Any]:
    directory = Path(session_dir)
    rows = _accepted_rows(directory)
    total = round(sum(float(row["net_g"]) for row in rows), 3)
    by_cultivar: dict[str, float] = defaultdict(float)
    for row in rows:
        by_cultivar[str(row["cultivar_normalized_name"])] += float(row["net_g"])
    source_hash = hashlib.sha256(("\n".join(canonical(row) for row in rows) + "\n").encode()).hexdigest()
    compiled_at = max([str(row["captured_at"]) for row in rows], default="1970-01-01T00:00:00Z")
    run_id = str(rows[0]["run_id"]) if rows else directory.name
    session_id = str(rows[0]["session_id"]) if rows else directory.name
    cultivator = str(rows[0].get("facility_id") or rows[0].get("cultivator") or "") if rows else ""
    operator_id = str(rows[0].get("operator_id") or "") if rows else ""
    stem = plain_export_stem(run_id, session_id)
    report: dict[str, Any] = {
        "report_id": f"harvest-report-{session_id}",
        "session_id": session_id,
        "run_id": run_id,
        "cultivator": cultivator,
        "operator_id": operator_id,
        "record_count": len(rows),
        "total_net_g": total,
        "cultivar_totals": {key: round(value, 3) for key, value in sorted(by_cultivar.items())},
        "records_sha256": source_hash,
        "compiled_at": compiled_at,
        "authoritative": False,
        "non_claims": list(NON_CLAIMS),
    }

    out = directory / "reports"
    out.mkdir(exist_ok=True)
    json_path = out / "harvest_run_report.json"
    summary_csv_path = out / "harvest_run_report.csv"
    plants_csv_path = out / f"{stem}_plants.csv"
    xlsx_path = out / f"{stem}_harvest.xlsx"
    docx_path = out / f"{stem}_harvest.docx"
    legacy_xlsx = out / "harvest_run_report.xlsx"
    legacy_docx = out / "harvest_run_report.docx"
    bundle_path = out / "handoff_bundle_manifest.json"

    atomic_json(json_path, report)

    # SR5: summary includes cultivator context; UTF-8; strain totals preserved.
    with summary_csv_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["cultivator", "strain", "net_g"])
        for cultivar, net in report["cultivar_totals"].items():
            writer.writerow(
                [
                    cultivator or _cultivator_for_strain(rows, cultivar),
                    cultivar,
                    net,
                ]
            )
        writer.writerow([cultivator, "TOTAL", total])

    with plants_csv_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(HEADERS)
        for row in rows:
            writer.writerow(row_for(row))

    _write_xlsx(xlsx_path, report, rows)
    _write_xlsx(legacy_xlsx, report, rows)

    _write_docx(docx_path, report, rows)
    _write_docx(legacy_docx, report, rows)

    artifacts = {
        "json": str(json_path.resolve()),
        "csv_summary": str(summary_csv_path.resolve()),
        "csv_plants": str(plants_csv_path.resolve()),
        "csv": str(plants_csv_path.resolve()),
        "xlsx": str(xlsx_path.resolve()),
        "docx": str(docx_path.resolve()),
        "xlsx_legacy": str(legacy_xlsx.resolve()),
        "docx_legacy": str(legacy_docx.resolve()),
        "handoff_bundle_manifest": str(bundle_path.resolve()),
    }
    report["artifacts"] = artifacts
    report["json_path"] = artifacts["json"]
    report["export_stem"] = stem

    bundle = {
        "manifest_type": "bbws.handoff_bundle",
        "version": "0.1.0",
        "series_id": "BBWS_SR5_run_artifact_polish",
        "session_id": session_id,
        "run_id": run_id,
        "cultivator": cultivator,
        "operator_id": operator_id,
        "record_count": len(rows),
        "total_net_g": total,
        "records_sha256": source_hash,
        "compiled_at": compiled_at,
        "authoritative": False,
        "non_claims": list(NON_CLAIMS),
        "artifacts": artifacts,
        "doctrine_cite": [
            "M:/SALVAGE/KNOWLEDGE/FOUNDATIONS/aos_book_spine_capsule_template_v0_1_10",
            "M:/SALVAGE/KNOWLEDGE/FOUNDATIONS/aos_output_response_meta_pack_v1_5_0",
        ],
    }
    atomic_json(bundle_path, bundle)
    report["handoff_bundle_manifest"] = artifacts["handoff_bundle_manifest"]
    return report


def reconcile_export_to_jsonl(session_dir: str | Path) -> dict[str, Any]:
    """Gate export derivatives against authoritative JSONL (counts, cultivar totals, SHA)."""
    directory = Path(session_dir)
    rows = _accepted_rows(directory)
    report = compile_report(directory)
    csv_path = Path(report["artifacts"]["csv_plants"])
    csv_rows = 0
    if csv_path.exists():
        with csv_path.open(encoding="utf-8", newline="") as handle:
            reader = csv.DictReader(handle)
            csv_rows = sum(1 for _ in reader)
    count_ok = csv_rows == len(rows)
    cultivar_ok = True
    by_cultivar: dict[str, float] = defaultdict(float)
    for row in rows:
        by_cultivar[str(row["cultivar_normalized_name"])] += float(row["net_g"])
    for key, value in report["cultivar_totals"].items():
        if round(by_cultivar.get(key, 0.0), 3) != round(float(value), 3):
            cultivar_ok = False
            break
    sha_ok = report["records_sha256"] == hashlib.sha256(
        ("\n".join(canonical(row) for row in rows) + "\n").encode()
    ).hexdigest()
    headers_ok = True
    if csv_path.exists():
        with csv_path.open(encoding="utf-8", newline="") as handle:
            reader = csv.reader(handle)
            header = next(reader, [])
            headers_ok = header == list(HEADERS) and "cultivator" in header and "strain" in header
    bundle_path = Path(report["artifacts"].get("handoff_bundle_manifest") or "")
    bundle_ok = bundle_path.exists()
    status = "pass" if count_ok and cultivar_ok and sha_ok and headers_ok and bundle_ok else "fail"
    receipt = {
        "gate": "export_jsonl_reconcile",
        "status": status,
        "session_id": report["session_id"],
        "jsonl_count": len(rows),
        "csv_plant_count": csv_rows,
        "count_ok": count_ok,
        "cultivar_totals_ok": cultivar_ok,
        "headers_ok": headers_ok,
        "bundle_ok": bundle_ok,
        "records_sha256": report["records_sha256"],
        "sha_ok": sha_ok,
        "authoritative": "records.jsonl",
        "handoff_bundle_manifest": str(bundle_path.resolve()) if bundle_ok else None,
        "non_claims": report["non_claims"],
    }
    out = directory / "reports" / "reconcile_receipt.json"
    atomic_json(out, receipt)
    receipt["receipt_path"] = str(out.resolve())
    return receipt
